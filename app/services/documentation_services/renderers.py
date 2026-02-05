"""
Strategy Pattern for Rendering DOCX content.
Generates a 'Business Overview' followed by 'Technical Reference' for each file.
Uses python-docx to generate rich documents with tables and formatting.
"""

from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

# IMPORTANT: Ensure your models.py is updated to support the nested structure!
from app.api.schemas.doc_models import FileSummary

class BaseRenderer(ABC):
    """Abstract Strategy."""

    def render_document(self, doc: Document, summary: FileSummary):
        """Template Method: Defines the standard document structure."""
        # 1. Main Title
        self._add_header(doc, f"Documentation: {summary.filename}", level=0)
        
        # 2. Business Overview (Standard for ALL files)
        self._render_business_section(doc, summary)
        
        doc.add_page_break()
        
        # 3. Technical Reference (Specific to file type)
        self._render_technical_section(doc, summary)

    def _render_business_section(self, doc: Document, summary: FileSummary):
        biz = summary.business_overview
        self._add_header(doc, "1.1 Purpose", level=2)
        # Handle cases where LLM might return purpose at top level or nested
        purpose = biz.get('purpose') or summary.business_purpose or "N/A"
        doc.add_paragraph(purpose)

    @abstractmethod
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        """Specific technical details (implemented by subclasses)."""
        pass

    def _add_header(self, doc: Document, text: str, level: int = 1):
        """Adds a styled heading."""
        doc.add_heading(text, level=level)

    def _add_labeled_paragraph(self, doc: Document, label: str, text: str):
        """Adds a line like 'Label: text' with bold label."""
        if not text: return
        p = doc.add_paragraph()
        runner = p.add_run(f"{label}: ")
        runner.bold = True
        p.add_run(str(text))

    def _add_bullet_list(self, doc: Document, items: List[str]):
        """Adds a bulleted list."""
        if not items: 
            doc.add_paragraph("None found.", style='List Bullet')
            return
        
        for item in items:
            # Handle potential dictionary in list (LLM artifact)
            text = item if isinstance(item, str) else str(item)
            doc.add_paragraph(text, style='List Bullet')

    def _create_key_value_table(self, doc: Document, title: str, data: Dict[str, Any]):
        """Creates a 2-column table for metadata."""
        if not data: return

        if title:
            self._add_header(doc, title, level=2)
            
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        self._set_cell_text(hdr_cells[0], "Attribute", bold=True, bg_color="E7E6E6")
        self._set_cell_text(hdr_cells[1], "Value", bold=True, bg_color="E7E6E6")

        for key, value in data.items():
            row_cells = table.add_row().cells
            # Format key (e.g., 'job_name' -> 'Job Name')
            formatted_key = key.replace('_', ' ').title()
            row_cells[0].text = formatted_key
            row_cells[1].text = str(value) if value else "-"
        
        doc.add_paragraph() # Spacer

    def _create_grid_table(self, doc: Document, title: str, data_list: List[Dict[str, Any]], headers: List[str]):
        """Creates a multi-column table from a list of dictionaries."""
        if not data_list: return

        if title:
            self._add_header(doc, title, level=2)
            
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            self._set_cell_text(hdr_cells[i], header, bold=True, bg_color="E7E6E6")

        # Data Rows
        for item in data_list:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                # Map Header "Step Name" -> key "step_name" (normalized)
                key = header.lower().replace(' ', '_')
                
                # Try explicit key, then normalized, then raw header
                val = item.get(key, item.get(header, "-"))
                
                # Handle nested dicts (rare but possible)
                if isinstance(val, dict):
                    val = str(val)
                    
                row_cells[i].text = str(val)

        doc.add_paragraph() # Spacer

    def _set_cell_text(self, cell, text: str, bold: bool = False, bg_color: str = None):
        """Helper to format cell text."""
        # Clear existing paragraph if needed (usually cells start with one empty p)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        if bg_color:
            self._set_cell_background(cell, bg_color)

    def _set_cell_background(self, cell, color_hex: str):
        """Sets cell background color (Shading) using OXML."""
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tc_pr.append(shd)

class CobolRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        self._add_header(doc, "2.1 Functional Capabilities", level=2)
        self._add_bullet_list(doc, tech.get('functional_capabilities', []))
        
        self._add_header(doc, "2.2 Key Operations", level=2)
        self._add_bullet_list(doc, tech.get('key_operations', []))
        
        # If there are data interactions (structured), make a table
        if tech.get('data_interactions') and isinstance(tech['data_interactions'][0], dict):
            self._create_grid_table(
                doc, "2.3 Data Interactions", 
                tech['data_interactions'], 
                ["Target", "Operation"]
            )
        else:
            # Fallback for older prompts
            self._add_header(doc, "2.3 Technical Notes", level=2)
            self._add_bullet_list(doc, tech.get('technical_notes', []))

class PliRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        self._add_header(doc, "2.1 Procedure Logic", level=2)
        self._add_bullet_list(doc, tech.get('functional_capabilities', []))
        
        self._add_header(doc, "2.2 I/O and Calls", level=2)
        self._add_bullet_list(doc, tech.get('key_operations', []))
        
        self._add_header(doc, "2.3 Memory & Pointers", level=2)
        self._add_bullet_list(doc, tech.get('technical_notes', []))

class AssemblyRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        self._add_header(doc, "2.1 Register Usage", level=2)
        self._add_bullet_list(doc, tech.get('register_usage', []))
        
        self._add_header(doc, "2.2 Macros & Services", level=2)
        self._add_bullet_list(doc, tech.get('key_operations', []))
        
        self._add_header(doc, "2.3 Logic Flow", level=2)
        self._add_bullet_list(doc, tech.get('functional_capabilities', []))

class JclRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        # 2.1 Identification Table
        if tech.get('job_header'):
            self._create_key_value_table(doc, "2.1 Job Identification", tech['job_header'])
            
        # 2.2 Steps Table
        steps = tech.get('steps', [])
        if steps and isinstance(steps[0], dict):
            self._create_grid_table(
                doc, "2.2 Execution Steps", steps, ["Step Name", "Program", "Description"]
            )
        else:
            self._add_bullet_list(doc, steps) # Fallback

        # 2.3 Datasets Table
        datasets = tech.get('io_datasets', [])
        if datasets and isinstance(datasets[0], dict):
            self._create_grid_table(
                doc, "2.3 IO Datasets", datasets, ["Dataset", "Usage"]
            )
        else:
            self._add_bullet_list(doc, datasets) # Fallback

class RexxRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        self._add_header(doc, "2.1 Automation Tasks", level=2)
        self._add_bullet_list(doc, tech.get('automation_tasks', []))
        
        self._add_header(doc, "2.2 External Utilities", level=2)
        self._add_bullet_list(doc, tech.get('external_utilities', []))
        
        self._add_header(doc, "2.3 Script Logic", level=2)
        self._add_bullet_list(doc, tech.get('technical_notes', []))

class Ca7Renderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        # 2.1 ID Table
        if tech.get('identification'):
            self._create_key_value_table(doc, "2.1 Schedule Identification", tech['identification'])
            
        # 2.2 Dependencies Table
        deps = tech.get('dependencies_triggers', [])
        if deps and isinstance(deps[0], dict):
            self._create_grid_table(
                doc, "2.2 Dependencies", deps, ["Type", "Target", "Condition"]
            )
        else:
            self._add_bullet_list(doc, deps)

        # 2.3 User Reqs
        self._add_header(doc, "2.3 User Requirements", level=2)
        self._add_bullet_list(doc, tech.get('user_requirements', []))

class SqlRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        # Header info
        self._add_labeled_paragraph(doc, "Entity Name", tech.get('table_name') or summary.table_name)
        
        # Key Fields List
        if tech.get('key_fields'):
            self._add_header(doc, "2.1 Key Fields", level=2)
            # Assuming list of dicts from new prompt
            fields = tech['key_fields']
            if fields and isinstance(fields[0], dict):
                self._create_grid_table(doc, None, fields, ["Field", "Description"])
            else:
                self._add_bullet_list(doc, fields)

        # Schema Table
        struct = tech.get('table_structure') or summary.table_structure
        if struct and isinstance(struct[0], dict):
            self._create_grid_table(
                doc, "2.2 Data Structure", struct, ["Column Name", "Type", "Nullable"]
            )
        elif struct:
            self._add_bullet_list(doc, struct)

class BindRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        self._add_header(doc, "2.1 Bind Configuration", level=2)
        self._add_bullet_list(doc, tech.get('key_parameters', []))
        
        self._add_header(doc, "2.2 Package List (DBRMs)", level=2)
        self._add_bullet_list(doc, tech.get('configuration_areas', []))

class ParmlibRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        # Parameters Table
        params = tech.get('key_parameters', [])
        if params and isinstance(params[0], dict):
            self._create_grid_table(
                doc, "2.1 System Parameters", params, ["Name", "Value", "Description"]
            )
        else:
            self._add_bullet_list(doc, params)
            
        self._add_header(doc, "2.2 Configured Areas", level=2)
        self._add_bullet_list(doc, tech.get('configuration_areas', []))

class FlatFileRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        
        tech = getattr(summary, 'technical_analysis', {})
        
        struct = tech.get('table_structure', [])
        if struct and isinstance(struct[0], dict):
            self._create_grid_table(
                doc, "2.1 File Layout", struct, ["Column Name", "Type", "Nullable"]
            )
        
        self._add_header(doc, "2.2 Volume & Characteristics", level=2)
        self._add_bullet_list(doc, tech.get('technical_notes', []))

class DefaultRenderer(BaseRenderer):
    def _render_technical_section(self, doc: Document, summary: FileSummary):
        self._add_header(doc, "2. Technical Reference", level=1)
        self._add_bullet_list(doc, summary.notes)

class RendererFactory:
    _RENDERERS = {
        'cobol': CobolRenderer(),
        'pli': PliRenderer(),
        'assembly': AssemblyRenderer(),
        'hlasm': AssemblyRenderer(),
        'rexx': RexxRenderer(),
        'jcl': JclRenderer(),
        'proc': JclRenderer(),
        'ca7': Ca7Renderer(),
        'dclgen': SqlRenderer(),
        'sql': SqlRenderer(),
        'copybook': SqlRenderer(),
        'pli_copybook': SqlRenderer(),
        'bind': BindRenderer(),
        'parmlib': ParmlibRenderer(),
        'csv': FlatFileRenderer(),
        'fixed_length': FlatFileRenderer(),
        'flat_file': FlatFileRenderer(),
    }

    @staticmethod
    def get_renderer(file_type: str) -> BaseRenderer:
        return RendererFactory._RENDERERS.get(file_type.lower(), DefaultRenderer())